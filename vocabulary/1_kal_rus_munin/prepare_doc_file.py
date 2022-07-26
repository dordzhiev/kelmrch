from docx import Document as Doc


def main():
    with open('erdni.sharmandzhiev@gmail.com-dictionary.docx', 'rb') as f:
        document = Doc(f)
    paragraphs = document.paragraphs
    new_document = Doc()

    for paragraph in paragraphs:
        strings = ['']
        bold_indexes_list = [[]]
        i = 0
        if '\n' in paragraph.text:
            lf_i = paragraph.text.index('\n')
            for run in paragraph.runs:
                if len(strings[i]) + len(run.text) < lf_i + 1:
                    strings[i] += run.text
                    if run.bold and run.text:
                        bold_indexes_list[i].append((len(strings[i]) - len(run.text), len(strings[i])))
                else:
                    segment_last_index = run.text.find('\n')
                    segment = run.text[:segment_last_index]
                    strings[i] += segment
                    if run.bold and segment:
                        bold_indexes_list[i].append((len(strings[i]) - len(segment), len(strings[i])))
                    strings.append('')
                    bold_indexes_list.append(list())
                    i += 1
                    segment = run.text.replace('\n', '')[segment_last_index:]
                    strings[i] += segment
                    if run.bold and segment:
                        bold_indexes_list[i].append((len(strings[i]) - len(segment), len(strings[i])))
                    lf_i = paragraph.text.index('\n', lf_i + 2) \
                        if '\n' in paragraph.text[lf_i + 2:] \
                        else len(paragraph.text)
        else:
            for run in paragraph.runs:
                strings[i] += run.text
                if run.bold and run.text:
                    bold_indexes_list[i].append((len(strings[i]) - len(run.text), len(strings[i])))
        for i, string in enumerate(strings):
            new_paragraph = new_document.add_paragraph()
            last_index = 0
            for bold in bold_indexes_list[i]:
                if bold[1] != 0:
                    new_paragraph.add_run(string[last_index:bold[0]])
                new_paragraph.add_run(string[bold[0]:bold[1]]).bold = True
                last_index = bold[1]
            new_paragraph.add_run(string[last_index:])
    with open('vocabulary/new_document.docx', 'wb') as f:
        new_document.save(f)


if __name__ == '__main__':
    main()
